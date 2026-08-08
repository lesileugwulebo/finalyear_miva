import re
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_formatted_runs(paragraph, text, base_font_name='Times New Roman', base_font_size=Pt(12)):
    # Simple markdown inline formatting parser (**bold**, *italic*, `code`)
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    tokens = pattern.split(text)
    
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        run.font.name = base_font_name
        run.font.size = base_font_size
        
        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x88, 0x15, 0x15)
        else:
            run.text = token

def parse_markdown_table(lines):
    table_data = []
    for line in lines:
        if not line.strip().startswith('|'):
            continue
        # Check if line is divider (|---|---|)
        if re.match(r'^\s*\|(?:\s*:?-+:?\s*\|)+\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        table_data.append(cells)
    return table_data

def build_docx():
    doc = Document()

    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Normal Style Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal_style.paragraph_format.line_spacing = 1.25
    normal_style.paragraph_format.space_after = Pt(6)

    # Read markdown files
    scratch_dir = Path(r"C:\Users\Anna\.gemini\antigravity\scratch")
    file1 = scratch_dir / "revised_chapter1_and_chapter2.md"
    file2 = scratch_dir / "revised_chapter3_to_chapter6.md"

    content = ""
    if file1.exists():
        content += file1.read_text(encoding="utf-8") + "\n\n"
    if file2.exists():
        content += file2.read_text(encoding="utf-8") + "\n\n"

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_text = []
    code_lang = ""
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                run = p.add_run("\n".join(code_block_text))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

                code_block_text = []
                in_code_block = False
            else:
                in_code_block = True
                code_lang = line.strip().lstrip('`').strip()
                code_block_text = []
            i += 1
            continue

        if in_code_block:
            code_block_text.append(line)
            i += 1
            continue

        # Table handling
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                # Render table
                table_data = parse_markdown_table(table_lines)
                if table_data:
                    rows_count = len(table_data)
                    cols_count = max(len(r) for r in table_data)
                    
                    t = doc.add_table(rows=rows_count, cols=cols_count)
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    t.autofit = False

                    for r_idx, row_cells in enumerate(table_data):
                        for c_idx, cell_value in enumerate(row_cells):
                            if c_idx < cols_count:
                                cell = t.cell(r_idx, c_idx)
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                                p = cell.paragraphs[0]
                                p.paragraph_format.space_after = Pt(2)
                                p.paragraph_format.space_before = Pt(2)
                                p.paragraph_format.line_spacing = 1.15

                                if r_idx == 0:
                                    set_cell_background(cell, "1F497D") # Header Dark Blue
                                    add_formatted_runs(p, cell_value, base_font_size=Pt(10.5))
                                    for r in p.runs:
                                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                        r.bold = True
                                else:
                                    if r_idx % 2 == 1:
                                        set_cell_background(cell, "F2F5F9") # Alternate zebra stripe
                                    else:
                                        set_cell_background(cell, "FFFFFF")
                                    add_formatted_runs(p, cell_value, base_font_size=Pt(10))

                in_table = False
                table_lines = []

        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[2:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) # Navy
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[3:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[4:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13.5)
            run.bold = True
            run.font.color.rgb = RGBColor(0x36, 0x60, 0x92)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[5:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(0x36, 0x60, 0x92)
        # Bullet list items
        elif re.match(r'^\s*[\-\*]\s+', line):
            indent_level = len(re.match(r'^\s*', line).group(0)) // 2
            text = re.sub(r'^\s*[\-\*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, text)
        # Numbered list items
        elif re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, text)
        # Horizontal Rule
        elif line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("―" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        # Paragraphs
        elif line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            add_formatted_runs(p, line.strip())

        i += 1

    output_path = scratch_dir / "final_year.docx"
    doc.save(str(output_path))
    print(f"Successfully generated Word document: {output_path}")

if __name__ == "__main__":
    build_docx()
