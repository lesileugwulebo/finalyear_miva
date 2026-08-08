import re
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_formatted_runs(paragraph, text, base_font_name='Times New Roman', base_font_size=Pt(12)):
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    tokens = pattern.split(text)
    
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        run.font.name = base_font_name
        run.font.size = base_font_size
        
        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x88, 0x15, 0x15)
        else:
            run.text = token

def parse_markdown_table(lines):
    table_data = []
    for line in lines:
        if not line.strip().startswith('|'):
            continue
        if re.match(r'^\s*\|(?:\s*:?-+:?\s*\|)+\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        table_data.append(cells)
    return table_data

def add_front_matter(doc):
    # Title Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(24)
    run_title = p_title.add_run("DESIGN, IMPLEMENTATION, AND EVALUATION OF A SECURE AWS–AZURE MULTI-CLOUD ARCHITECTURE FOR ENTERPRISE WORKLOADS USING ZERO TRUST AND INFRASTRUCTURE AS CODE\n")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    r_sub = p_sub.add_run("A Professional Master's Project Report submitted to the Department of Information Technology, Miva Open University, in partial fulfilment of the requirements for the award of the degree of\n\nMASTER OF INFORMATION TECHNOLOGY (MIT)")
    r_sub.font.name = 'Times New Roman'
    r_sub.font.size = Pt(12)

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_before = Pt(48)
    p_author.paragraph_format.space_after = Pt(48)
    r_auth = p_author.add_run("BY\n\nLESILE UGWULEBO\nMATRICULATION NO: MIT/2025/00142\n\nDEPARTMENT OF INFORMATION TECHNOLOGY\nMIVA OPEN UNIVERSITY, NIGERIA\n\nAUGUST 2026")
    r_auth.font.name = 'Times New Roman'
    r_auth.font.size = Pt(12)
    r_auth.bold = True

    doc.add_page_break()

    # Declaration Page
    p_dec_h = doc.add_paragraph()
    p_dec_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_dec_h = p_dec_h.add_run("DECLARATION\n")
    r_dec_h.font.name = 'Times New Roman'
    r_dec_h.font.size = Pt(14)
    r_dec_h.bold = True

    p_dec_body = doc.add_paragraph()
    p_dec_body.paragraph_format.line_spacing = 1.25
    p_dec_body.paragraph_format.space_after = Pt(18)
    add_formatted_runs(p_dec_body, "I, Lesile Ugwulebo, hereby declare that this Professional Master’s Project Report titled **'DESIGN, IMPLEMENTATION, AND EVALUATION OF A SECURE AWS–AZURE MULTI-CLOUD ARCHITECTURE FOR ENTERPRISE WORKLOADS USING ZERO TRUST AND INFRASTRUCTURE AS CODE'** is an original account of my research work conducted under the supervision of the Department of Information Technology, Miva Open University.\n\nThis work has not been submitted previously in part or in full for any other degree, diploma, or professional qualification at this or any other academic institution. All literature sources, tools, frameworks, and secondary data utilised have been duly acknowledged through appropriate citations and referencing.")

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(36)
    r_sig = p_sig.add_run("_________________________\t\t\t_________________________\nLesile Ugwulebo\t\t\t\t\tDate")
    r_sig.font.name = 'Times New Roman'

    doc.add_page_break()

    # Certification Page
    p_cert_h = doc.add_paragraph()
    p_cert_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cert_h = p_cert_h.add_run("CERTIFICATION AND APPROVAL PAGE\n")
    r_cert_h.font.name = 'Times New Roman'
    r_cert_h.font.size = Pt(14)
    r_cert_h.bold = True

    p_cert_body = doc.add_paragraph()
    p_cert_body.paragraph_format.line_spacing = 1.25
    p_cert_body.paragraph_format.space_after = Pt(24)
    add_formatted_runs(p_cert_body, "This is to certify that this Professional Master’s Project Report by **Lesile Ugwulebo (Matriculation No: MIT/2025/00142)** has been examined and approved as meeting the requirements of the Department of Information Technology, Miva Open University, for the award of the degree of Master of Information Technology (MIT).")

    p_sigs = doc.add_paragraph()
    p_sigs.paragraph_format.space_before = Pt(24)
    p_sigs.paragraph_format.line_spacing = 1.5
    r_sigs = p_sigs.add_run("_____________________________________\t\t\t___________________\nProject Supervisor\t\t\t\t\t\t\tDate\n\n\n_____________________________________\t\t\t___________________\nHead of Department\t\t\t\t\t\t\tDate\n\n\n_____________________________________\t\t\t___________________\nExternal Examiner\t\t\t\t\t\t\tDate")
    r_sigs.font.name = 'Times New Roman'

    doc.add_page_break()

    # Acknowledgements Page
    p_ack_h = doc.add_paragraph()
    p_ack_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ack_h = p_ack_h.add_run("ACKNOWLEDGEMENTS\n")
    r_ack_h.font.name = 'Times New Roman'
    r_ack_h.font.size = Pt(14)
    r_ack_h.bold = True

    p_ack_body = doc.add_paragraph()
    p_ack_body.paragraph_format.line_spacing = 1.25
    add_formatted_runs(p_ack_body, "First and foremost, I express my sincere gratitude to Almighty God for guidance, wisdom, and health throughout the duration of this Master of Information Technology (MIT) programme.\n\nMy deepest appreciation goes to my project supervisor for invaluable academic guidance, constructive critique, and continuous encouragement. Special thanks to the faculty and management of Miva Open University for establishing an excellent, solution-driven professional learning environment.\n\nFinally, I extend my heartfelt thanks to my family, colleagues, and peers for their continuous moral support and understanding during the execution of this project.")

    doc.add_page_break()

    # Abstract Page
    p_abs_h = doc.add_paragraph()
    p_abs_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_abs_h = p_abs_h.add_run("ABSTRACT\n")
    r_abs_h.font.name = 'Times New Roman'
    r_abs_h.font.size = Pt(14)
    r_abs_h.bold = True

    p_abs_body = doc.add_paragraph()
    p_abs_body.paragraph_format.line_spacing = 1.25
    add_formatted_runs(p_abs_body, "Modern enterprise information technology increasingly relies on multi-cloud strategies connecting Amazon Web Services (AWS) and Microsoft Azure to optimize service availability, access differentiated features, and mitigate single-vendor concentration risk. However, uncoordinated multi-cloud connectivity frequently introduces significant security and operational vulnerabilities, including inconsistent firewall policy enforcement, unmonitored attack paths for lateral movement, fragmented identity administration, disparate logging formats, and unverified failover resilience. Furthermore, in the Nigerian enterprise context, ungoverned inter-cloud data flows make it difficult to demonstrate compliance with the National Cloud Policy 2025 and the Nigeria Data Protection Act (NDPA) 2023. This Professional Master’s Project addresses these challenges through the design, implementation, and empirical evaluation of a secure, practice-oriented AWS–Azure reference architecture using Zero Trust principles, defence-in-depth controls, and modular Infrastructure as Code (IaC).\n\nFollowing Design Science Research Methodology (DSRM), the project engineered a 5-plane hub-and-spoke multi-cloud architecture utilizing AWS Transit Gateway (TGW) and Azure Active-Active VPN Gateway connected via route-based IKEv2 IPsec VPN with Border Gateway Protocol (BGP-4). Non-overlapping IPv4 address spaces (AWS 10.10.0.0/16 and Azure 10.20.0.0/16) were established with strict tier-level micro-segmentation. Workforce identity was centralized by federating Microsoft Entra ID with AWS IAM Identity Center using SAML 2.0 and automated SCIM provisioning. The entire multi-cloud infrastructure was codified using modular Terraform (HCL) with encrypted remote state management.\n\nEmpirical evaluation validated the system's security, performance, and resilience. Security posture assessment using Prowler and ScoutSuite confirmed 0 Critical and 0 High findings within the project scope, verifying that direct web-to-database and cross-cloud database bypass paths were blocked. Active network benchmarking demonstrated an average inter-cloud round-trip latency of 35.15 ms—well within the sub-100ms threshold—and TCP throughput scaling up to 498.9 Mbps over parallel streams. Simulated fault injection (disabling an active IPsec tunnel during a continuous 1-second probe session) demonstrated dynamic BGP route convergence with a measured Recovery Time Objective (RTO) of 4.1 seconds. The project contributes a deployable reference blueprint, provider control mappings, IaC security governance practices, and a phased adoption roadmap for Nigerian enterprises.\n\n**Keywords**: Multi-Cloud Architecture, Zero Trust, Amazon Web Services, Microsoft Azure, BGP IPsec VPN, Infrastructure as Code, Microsoft Entra ID, Nigeria Data Protection Act.")

    doc.add_page_break()

    # Abbreviations & Acronyms Page
    p_abb_h = doc.add_paragraph()
    p_abb_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_abb_h = p_abb_h.add_run("LIST OF ABBREVIATIONS AND ACRONYMS\n")
    r_abb_h.font.name = 'Times New Roman'
    r_abb_h.font.size = Pt(14)
    r_abb_h.bold = True

    abbreviations = [
        ("ALB", "Application Load Balancer"),
        ("API", "Application Programming Interface"),
        ("ASN", "Autonomous System Number"),
        ("AWS", "Amazon Web Services"),
        ("BGP", "Border Gateway Protocol"),
        ("CCM", "Cloud Controls Matrix (Cloud Security Alliance)"),
        ("CIDR", "Classless Inter-Domain Routing"),
        ("CLI", "Command Line Interface"),
        ("CSA", "Cloud Security Alliance"),
        ("DSRM", "Design Science Research Methodology"),
        ("EBS", "Elastic Block Store"),
        ("EC2", "Elastic Compute Cloud"),
        ("ESP", "Encapsulating Security Payload"),
        ("GAID", "General Application and Implementation Directive (NDPC)"),
        ("GCP", "Google Cloud Platform"),
        ("HCL", "HashiCorp Configuration Language"),
        ("IaaS", "Infrastructure as a Service"),
        ("IaC", "Infrastructure as Code"),
        ("IdP", "Identity Provider"),
        ("IKE", "Internet Key Exchange"),
        ("IPsec", "Internet Protocol Security"),
        ("KMS", "Key Management Service"),
        ("MCN", "Multi-Cloud Networking"),
        ("MFA", "Multi-Factor Authentication"),
        ("MIT", "Master of Information Technology"),
        ("NACL", "Network Access Control List"),
        ("NDPA", "Nigeria Data Protection Act 2023"),
        ("NDPC", "Nigeria Data Protection Commission"),
        ("NIST", "National Institute of Standards and Technology"),
        ("NITDA", "National Information Technology Development Agency"),
        ("NSG", "Network Security Group"),
        ("NUC", "National Universities Commission"),
        ("PaaS", "Platform as a Service"),
        ("PEP", "Policy Enforcement Point"),
        ("RBAC", "Role-Based Access Control"),
        ("RTO", "Recovery Time Objective"),
        ("RTT", "Round-Trip Time"),
        ("SAML", "Security Assertion Markup Language"),
        ("SaaS", "Software as a Service"),
        ("SCIM", "System for Cross-domain Identity Management"),
        ("SIEM", "Security Information and Event Management"),
        ("SP", "Service Provider"),
        ("SSO", "Single Sign-On"),
        ("STS", "Security Token Service"),
        ("TGW", "AWS Transit Gateway"),
        ("UDR", "User-Defined Route"),
        ("VMC", "Virtual Machine Scale"),
        ("VNG", "Azure Virtual Network Gateway"),
        ("VNet", "Azure Virtual Network"),
        ("VPC", "Amazon Virtual Private Cloud"),
        ("VPN", "Virtual Private Network"),
        ("ZTA", "Zero Trust Architecture")
    ]

    t_abb = doc.add_table(rows=len(abbreviations)+1, cols=2)
    t_abb.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    cell_00 = t_abb.cell(0, 0)
    cell_01 = t_abb.cell(0, 1)
    set_cell_background(cell_00, "1F497D")
    set_cell_background(cell_01, "1F497D")
    p00 = cell_00.paragraphs[0]
    p01 = cell_01.paragraphs[0]
    r00 = p00.add_run("Abbreviation / Acronym")
    r01 = p01.add_run("Full Meaning")
    for r in [r00, r01]:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (abbr, meaning) in enumerate(abbreviations, start=1):
        c0 = t_abb.cell(idx, 0)
        c1 = t_abb.cell(idx, 1)
        if idx % 2 == 1:
            set_cell_background(c0, "F2F5F9")
            set_cell_background(c1, "F2F5F9")
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        
        p0 = c0.paragraphs[0]
        p1 = c1.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        
        r0 = p0.add_run(abbr)
        r1 = p1.add_run(meaning)
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(10)
        r0.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(10)

    doc.add_page_break()

# Mapping of figure titles to generated PNG diagram images
FIGURE_IMAGE_MAP = {
    "figure 1.1": "fig1_1_conceptual_framework.png",
    "figure 2.1": "fig2_1_regulatory_framework.png",
    "figure 2.2": "fig2_2_intercloud_topology.png",
    "figure 2.3": "fig2_3_zerotrust_planes.png",
    "figure 2.4": "fig2_4_saml_federation.png",
    "figure 2.5": "fig2_5_solution_quadrant.png",
    "figure 3.1": "fig3_1_dsrm_cycle.png",
    "figure 3.2": "fig3_2_five_plane_architecture.png",
    "figure 5.1": "fig5_1_traffic_segmentation.png",
    "figure 5.2": "fig5_2_prowler_findings.png",
    "figure 5.3": "fig5_3_latency_chart.png",
    "figure 5.4": "fig5_4_throughput_chart.png",
    "figure 5.5": "fig5_5_failover_sequence.png",
    "figure 6.1": "fig6_1_adoption_roadmap.png",
}

def build_docx():
    doc = Document()

    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Add Front Matter
    add_front_matter(doc)

    # Base Normal Style Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal_style.paragraph_format.line_spacing = 1.25
    normal_style.paragraph_format.space_after = Pt(6)

    scratch_dir = Path(r"C:\Users\Anna\.gemini\antigravity\scratch")
    diagrams_dir = scratch_dir / "diagrams"
    
    file1 = scratch_dir / "revised_chapter1_and_chapter2.md"
    file2 = scratch_dir / "revised_chapter3_to_chapter6.md"

    content = ""
    if file1.exists():
        content += file1.read_text(encoding="utf-8") + "\n\n"
    if file2.exists():
        content += file2.read_text(encoding="utf-8") + "\n\n"

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_text = []
    code_lang = ""
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # If code block was mermaid, skip outputting raw text since diagram PNG will be inserted by figure caption
                if code_lang.lower() == 'mermaid':
                    pass
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    p.paragraph_format.right_indent = Inches(0.4)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    
                    run = p.add_run("\n".join(code_block_text))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

                code_block_text = []
                in_code_block = False
            else:
                in_code_block = True
                code_lang = line.strip().lstrip('`').strip()
                code_block_text = []
            i += 1
            continue

        if in_code_block:
            code_block_text.append(line)
            i += 1
            continue

        # Table handling
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                table_data = parse_markdown_table(table_lines)
                if table_data:
                    rows_count = len(table_data)
                    cols_count = max(len(r) for r in table_data)
                    
                    t = doc.add_table(rows=rows_count, cols=cols_count)
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    t.autofit = False

                    for r_idx, row_cells in enumerate(table_data):
                        for c_idx, cell_value in enumerate(row_cells):
                            if c_idx < cols_count:
                                cell = t.cell(r_idx, c_idx)
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                                p = cell.paragraphs[0]
                                p.paragraph_format.space_after = Pt(2)
                                p.paragraph_format.space_before = Pt(2)
                                p.paragraph_format.line_spacing = 1.15

                                if r_idx == 0:
                                    set_cell_background(cell, "1F497D")
                                    add_formatted_runs(p, cell_value, base_font_size=Pt(10.5))
                                    for r in p.runs:
                                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                        r.bold = True
                                else:
                                    if r_idx % 2 == 1:
                                        set_cell_background(cell, "F2F5F9")
                                    else:
                                        set_cell_background(cell, "FFFFFF")
                                    add_formatted_runs(p, cell_value, base_font_size=Pt(10))

                in_table = False
                table_lines = []

        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[2:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[3:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(15)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[4:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = RGBColor(0x36, 0x60, 0x92)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[5:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(0x36, 0x60, 0x92)
        elif re.match(r'^\s*[\-\*]\s+', line):
            indent_level = len(re.match(r'^\s*', line).group(0)) // 2
            text = re.sub(r'^\s*[\-\*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, text)
        elif re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, text)
        elif line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("―" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        elif line.strip():
            # Check if line contains a Figure caption like "Figure 1.1: ..."
            fig_match = re.search(r'\b(Figure\s+\d+\.\d+)', line, re.IGNORECASE)
            if fig_match:
                fig_key = fig_match.group(1).lower()
                if fig_key in FIGURE_IMAGE_MAP:
                    img_filename = FIGURE_IMAGE_MAP[fig_key]
                    img_path = diagrams_dir / img_filename
                    if img_path.exists():
                        # Insert Diagram PNG Image
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(12)
                        p_img.paragraph_format.space_after = Pt(4)
                        p_img.paragraph_format.keep_with_next = True
                        p_img.add_run().add_picture(str(img_path), width=Inches(6.0))

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            if fig_match:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_runs(p, line.strip())

        i += 1

    output_path = scratch_dir / "final_year.docx"
    doc.save(str(output_path))
    print(f"Successfully generated Word Document with embedded high-res PNG diagrams: {output_path}")

if __name__ == "__main__":
    build_docx()
